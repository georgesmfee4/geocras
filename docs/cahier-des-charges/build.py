# -*- coding: utf-8 -*-
"""Compose le cahier des charges : source balisée → DOCX → PDF.

Deux passes : la première mesure la position réelle de chaque titre dans le PDF,
la seconde reconstruit le document avec un sommaire dont les numéros de page
sont justes.
"""
import os, re, subprocess, sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Emu

HERE = os.path.dirname(os.path.abspath(__file__))

SERIF = "IBM Plex Serif"
SANS = "IBM Plex Sans"
MONO = "IBM Plex Mono"

INK = RGBColor(0x1C, 0x1A, 0x17)
INK2 = RGBColor(0x55, 0x51, 0x4A)
MUTED = RGBColor(0x8A, 0x85, 0x78)
RED = RGBColor(0xC6, 0x2A, 0x26)
GREEN = RGBColor(0x24, 0x6B, 0x45)

PAGE_W = Cm(21.0)
MARGIN_L = Cm(2.4)
MARGIN_R = Cm(2.2)
CONTENT_W = PAGE_W - MARGIN_L - MARGIN_R


# --------------------------------------------------------------------------- #
# helpers XML
# --------------------------------------------------------------------------- #
def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn("w:" + k), v)
    return e


def shade(par_or_cell, hexcolor):
    pr = par_or_cell._element.get_or_add_tcPr() if hasattr(par_or_cell._element, "get_or_add_tcPr") \
        else par_or_cell._p.get_or_add_pPr()
    pr.append(_el("w:shd", val="clear", color="auto", fill=hexcolor))


def border(par, edge="bottom", size=6, color="C62A26", space=4):
    pPr = par._p.get_or_add_pPr()
    pbdr = pPr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        pPr.append(pbdr)
    e = _el("w:" + edge, val="single", sz=str(size), space=str(space), color=color)
    pbdr.append(e)


def keep_with_next(par):
    par._p.get_or_add_pPr().append(_el("w:keepNext", val="1"))


def page_break_before(par):
    par._p.get_or_add_pPr().append(_el("w:pageBreakBefore", val="1"))


def field(par, instr, placeholder="1"):
    r = par.add_run()
    b = _el("w:fldChar", fldCharType="begin"); r._r.append(b)
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    r._r.append(it)
    r._r.append(_el("w:fldChar", fldCharType="separate"))
    t = OxmlElement("w:t"); t.text = placeholder; r._r.append(t)
    r._r.append(_el("w:fldChar", fldCharType="end"))
    return r


def style_run(run, font=SERIF, size=10.5, color=INK, bold=False, italic=False, caps=False, spacing=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    rPr = run._r.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rPr.append(rf)
    for a in ("ascii", "hAnsi", "cs", "eastAsia"):
        rf.set(qn("w:" + a), font)
    if caps:
        rPr.append(_el("w:caps", val="1"))
    if spacing:
        rPr.append(_el("w:spacing", val=str(int(spacing * 20))))
    return run


# --------------------------------------------------------------------------- #
# parsing de la source
# --------------------------------------------------------------------------- #
def parse(path):
    blocks, i = [], 0
    lines = open(path, encoding="utf-8").read().split("\n")
    while i < len(lines):
        ln = lines[i]
        st = ln.strip()
        if not st:
            i += 1; continue
        m = re.match(r"^#(\d)\s+(.*)$", st)
        if m:
            blocks.append(("h" + m.group(1), m.group(2))); i += 1; continue
        if st.startswith("[FIG2]"):
            parts = [x.strip() for x in st[6:].split("|")]
            blocks.append(("fig2", (parts[0], parts[1], parts[2] if len(parts) > 2 else "",
                                    float(parts[3]) if len(parts) > 3 else 9.0)))
            i += 1; continue
        if st.startswith("[FIG]"):
            parts = [x.strip() for x in st[5:].split("|")]
            blocks.append(("fig", (parts[0], parts[1] if len(parts) > 1 else "",
                                   float(parts[2]) if len(parts) > 2 else 20.0)))
            i += 1; continue
        if st == "[PB]":
            blocks.append(("pb", None)); i += 1; continue
        if st.startswith("[TAB]"):
            cap = st[5:].strip(); rows = []; i += 1
            while lines[i].strip() != "[/TAB]":
                rows.append([c.strip() for c in lines[i].split("|")]); i += 1
            blocks.append(("tab", (cap, rows))); i += 1; continue
        if st.startswith("[CODE]"):
            cap = st[6:].strip(); body = []; i += 1
            while lines[i].strip() != "[/CODE]":
                body.append(lines[i]); i += 1
            blocks.append(("code", (cap, body))); i += 1; continue
        if st.startswith("[NOTE]"):
            title = st[6:].strip(); body = []; i += 1
            while lines[i].strip() != "[/NOTE]":
                body.append(lines[i].strip()); i += 1
            blocks.append(("note", (title, " ".join(body)))); i += 1; continue
        if st.startswith("[LIST]"):
            items = []; i += 1
            while lines[i].strip() != "[/LIST]":
                if lines[i].strip():
                    items.append(lines[i].strip().lstrip("- ").strip())
                i += 1
            blocks.append(("list", items)); i += 1; continue
        if st.startswith("[COVER]"):
            body = []; i += 1
            while lines[i].strip() != "[/COVER]":
                body.append(lines[i].rstrip()); i += 1
            blocks.append(("cover", body)); i += 1; continue
        if st.startswith("[TOC]"):
            blocks.append(("toc", None)); i += 1; continue
        if st.startswith("[FIGTOC]"):
            blocks.append(("figtoc", None)); i += 1; continue
        # paragraphe : lignes consécutives jusqu'à une ligne vide
        buf = [st]; i += 1
        while i < len(lines) and lines[i].strip() and not re.match(r"^(\[|#\d)", lines[i].strip()):
            buf.append(lines[i].strip()); i += 1
        blocks.append(("p", " ".join(buf)))
    return blocks


RICH = re.compile(r"(\*\*.+?\*\*|`.+?`|__.+?__)")


def add_rich(par, text, size=10.5, font=SERIF, color=INK):
    """Gras **…**, mono `…`, italique __…__."""
    for part in RICH.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            style_run(par.add_run(part[2:-2]), font=font, size=size, color=color, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            style_run(par.add_run(part[1:-1]), font=MONO, size=size - 1.1, color=INK2)
        elif part.startswith("__") and part.endswith("__"):
            style_run(par.add_run(part[2:-2]), font=font, size=size, color=color, italic=True)
        else:
            style_run(par.add_run(part), font=font, size=size, color=color)


# --------------------------------------------------------------------------- #
# rendu
# --------------------------------------------------------------------------- #
class Builder:
    def __init__(self, toc_pages=None, fig_pages=None):
        self.doc = Document()
        self.toc_pages = toc_pages or {}
        self.fig_pages = fig_pages or {}
        self.headings = []          # (niveau, numéro, titre)
        self.figures = []           # (numéro, légende)
        self.counters = [0, 0, 0]
        self.fig_n = 0
        self.tab_n = 0
        self._setup()

    def _setup(self):
        sec = self.doc.sections[0]
        sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
        sec.top_margin, sec.bottom_margin = Cm(2.4), Cm(2.2)
        sec.left_margin, sec.right_margin = MARGIN_L, MARGIN_R
        sec.header_distance, sec.footer_distance = Cm(1.2), Cm(1.2)
        st = self.doc.styles["Normal"]
        st.font.name = SERIF
        st.font.size = Pt(10.5)
        st.paragraph_format.space_after = Pt(0)

    # -- primitives ------------------------------------------------------- #
    def para(self, text="", size=10.5, font=SERIF, color=INK, align="just",
             before=0, after=6, indent=0, line=1.3, rich=True):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = {"just": WD_ALIGN_PARAGRAPH.JUSTIFY, "left": WD_ALIGN_PARAGRAPH.LEFT,
                        "center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
        pf.space_before, pf.space_after = Pt(before), Pt(after)
        pf.line_spacing = line
        if indent:
            pf.left_indent = Cm(indent)
        if text:
            if rich:
                add_rich(p, text, size=size, font=font, color=color)
            else:
                style_run(p.add_run(text), font=font, size=size, color=color)
        return p

    def heading(self, level, title):
        n = int(level)
        if n == 0:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            style_run(p.add_run(title), font=SANS, size=16, color=INK, bold=True)
            border(p, "bottom", size=8, color="C62A26", space=6)
            keep_with_next(p)
            return p
        self.counters[n - 1] += 1
        for k in range(n, 3):
            self.counters[k] = 0
        num = ".".join(str(c) for c in self.counters[:n] if True)
        num = ".".join(str(self.counters[k]) for k in range(n))
        if n == 1:
            p = self.doc.add_paragraph()
            page_break_before(p)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            style_run(p.add_run("CHAPITRE " + str(self.counters[0])), font=SANS, size=9,
                      color=RED, bold=True, spacing=1.6)
            p2 = self.doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(16)
            style_run(p2.add_run(title), font=SANS, size=20, color=INK, bold=True)
            border(p2, "bottom", size=10, color="C62A26", space=8)
            keep_with_next(p2)
            return p2
        size = 13.5 if n == 2 else 11.5
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(15 if n == 2 else 11)
        p.paragraph_format.space_after = Pt(5)
        style_run(p.add_run(num + "   "), font=SANS, size=size, color=RED, bold=True)
        style_run(p.add_run(title), font=SANS, size=size, color=INK, bold=True)
        keep_with_next(p)
        return p

    def figure(self, src, caption, max_h=20.0):
        self.fig_n += 1
        path = os.path.join(HERE, src)
        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run()
        run.add_picture(path, width=self.fig_width(path, max_h))
        c = self.doc.add_paragraph()
        c.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_after = Pt(12)
        style_run(c.add_run("Figure %d — " % self.fig_n), font=SANS, size=9, color=RED, bold=True)
        style_run(c.add_run(caption), font=SANS, size=9, color=INK2)
        return p

    def fig_width(self, path, max_h_cm=20.0):
        from PIL import Image
        w, h = Image.open(path).size
        max_w = CONTENT_W
        max_h = Cm(max_h_cm)
        ratio = h / w
        width = max_w
        if Emu(int(width * ratio)) > max_h:
            width = Emu(int(max_h / ratio))
        return width

    def figure_pair(self, src_a, src_b, caption, height_cm=9.0):
        """Deux captures d'écran alignées, avec une légende commune."""
        from PIL import Image
        self.fig_n += 1
        t = self.doc.add_table(rows=1, cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for cell, src in zip(t.rows[0].cells, (src_a, src_b)):
            path = os.path.join(HERE, src)
            w, h = Image.open(path).size
            par = cell.paragraphs[0]
            par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            par.paragraph_format.space_before = Pt(8)
            par.paragraph_format.space_after = Pt(2)
            par.add_run().add_picture(path, width=Emu(int(Cm(height_cm) * w / h)))
        c = self.doc.add_paragraph()
        c.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_after = Pt(12)
        style_run(c.add_run("Figure %d — " % self.fig_n), font=SANS, size=9, color=RED, bold=True)
        style_run(c.add_run(caption), font=SANS, size=9, color=INK2)
        return t

    def table(self, caption, rows):
        self.tab_n += 1
        if caption:
            c = self.doc.add_paragraph()
            c.paragraph_format.space_before = Pt(10)
            c.paragraph_format.space_after = Pt(4)
            style_run(c.add_run("Tableau %d — " % self.tab_n), font=SANS, size=9, color=RED, bold=True)
            style_run(c.add_run(caption), font=SANS, size=9, color=INK2)
            keep_with_next(c)
        ncols = max(len(r) for r in rows)
        t = self.doc.add_table(rows=0, cols=ncols)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = True
        for ri, row in enumerate(rows):
            cells = t.add_row().cells
            for ci in range(ncols):
                txt = row[ci] if ci < len(row) else ""
                cell = cells[ci]
                cell.text = ""
                par = cell.paragraphs[0]
                par.paragraph_format.space_before = Pt(3)
                par.paragraph_format.space_after = Pt(3)
                par.paragraph_format.line_spacing = 1.12
                if ri == 0:
                    add_rich(par, txt, size=9, font=SANS)
                    for r in par.runs:
                        r.bold = True
                    shade(cell, "F1EEE7")
                else:
                    add_rich(par, txt, size=9.2, font=SERIF)
        self._table_borders(t)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(8)
        return t

    def _table_borders(self, t):
        tbl = t._tbl
        pr = tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            col = "C9C3B6" if edge in ("insideH", "insideV") else "8A8578"
            sz = "4" if edge in ("insideH", "insideV") else "8"
            borders.append(_el("w:" + edge, val="single", sz=sz, space="0", color=col))
        pr.append(borders)

    def code(self, caption, body):
        if caption:
            c = self.doc.add_paragraph()
            c.paragraph_format.space_before = Pt(9)
            c.paragraph_format.space_after = Pt(3)
            style_run(c.add_run(caption), font=SANS, size=9, color=MUTED, italic=True)
            keep_with_next(c)
        for i, ln in enumerate(body):
            p = self.doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_before = Pt(4 if i == 0 else 0)
            pf.space_after = Pt(4 if i == len(body) - 1 else 0)
            pf.line_spacing = 1.12
            pf.left_indent = Cm(0.5)
            style_run(p.add_run(ln if ln.strip() else " "), font=MONO, size=8.6, color=INK)
            shade(p, "F6F4EF")
        self.doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def note(self, title, body):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(10); pf.space_after = Pt(2)
        pf.left_indent = Cm(0.4); pf.right_indent = Cm(0.2)
        style_run(p.add_run(title), font=SANS, size=9.5, color=RED, bold=True)
        border(p, "left", size=18, color="C62A26", space=8)
        shade(p, "FBF6F4")
        b = self.doc.add_paragraph()
        bf = b.paragraph_format
        bf.space_before = Pt(0); bf.space_after = Pt(11)
        bf.left_indent = Cm(0.4); bf.right_indent = Cm(0.2)
        bf.line_spacing = 1.28
        bf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_rich(b, body, size=9.8, font=SERIF, color=INK2)
        border(b, "left", size=18, color="C62A26", space=8)
        shade(b, "FBF6F4")

    def bullets(self, items):
        for it in items:
            p = self.doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Cm(0.75); pf.first_line_indent = Cm(-0.35)
            pf.space_after = Pt(4); pf.line_spacing = 1.28
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            style_run(p.add_run("—   "), font=SANS, size=10.5, color=RED, bold=True)
            add_rich(p, it, size=10.5)

    # -- pages liminaires -------------------------------------------------- #
    def cover(self, lines):
        for raw in lines:
            st = raw.strip()
            if st == "~":
                self.para("", after=10); continue
            if st.startswith("!"):
                self.para(st[1:].strip(), font=SANS, size=27, bold_hack=True, align="center",
                          after=4) if False else None
                p = self.doc.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(6)
                style_run(p.add_run(st[1:].strip()), font=SANS, size=30, color=INK,
                          bold=True, spacing=2.2)
                continue
            if st.startswith(">"):
                p = self.doc.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(4)
                style_run(p.add_run(st[1:].strip()), font=SANS, size=9.5, color=RED,
                          bold=True, spacing=2.4)
                continue
            if st.startswith("="):
                p = self.doc.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(12)
                style_run(p.add_run(st[1:].strip()), font=SANS, size=14.5, color=INK2)
                continue
            if st.startswith("-"):
                p = self.doc.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(3)
                style_run(p.add_run(st[1:].strip()), font=SERIF, size=11, color=INK2)
                continue
            if st == "___":
                p = self.doc.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(14)
                style_run(p.add_run(" " * 24), font=SANS, size=6)
                border(p, "bottom", size=12, color="C62A26", space=2)
                continue
            self.para(st, align="center", after=4)

    def toc_placeholder(self):
        self._toc_marker = len(self.doc.paragraphs)

    def render_toc(self):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(14)
        style_run(p.add_run("Sommaire"), font=SANS, size=20, color=INK, bold=True)
        border(p, "bottom", size=10, color="C62A26", space=8)
        for lvl, num, title in self.headings:
            key = (num, title)
            page = self.toc_pages.get(key, "")
            par = self.doc.add_paragraph()
            pf = par.paragraph_format
            pf.space_after = Pt(3 if lvl > 1 else 7)
            pf.space_before = Pt(9 if lvl <= 1 else 0)
            pf.left_indent = Cm({0: 0, 1: 0, 2: 0.65, 3: 1.4}[lvl])
            pf.tab_stops.add_tab_stop(CONTENT_W, WD_TAB_ALIGNMENT.RIGHT, 2)  # 2 = dotted
            size = 11.5 if lvl <= 1 else (10 if lvl == 2 else 9.4)
            if num:
                style_run(par.add_run(num + "   "), font=SANS, size=size,
                          color=RED if lvl == 1 else INK2, bold=(lvl == 1))
            style_run(par.add_run(title), font=SANS if lvl <= 1 else SERIF, size=size,
                      color=INK if lvl < 3 else INK2, bold=(lvl <= 1))
            style_run(par.add_run("\t"), font=SERIF, size=size)
            style_run(par.add_run(str(page)), font=MONO, size=size - 0.6,
                      color=INK if lvl == 1 else INK2, bold=(lvl == 1))

    def render_figtoc(self):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        style_run(p.add_run("Table des figures"), font=SANS, size=15, color=INK, bold=True)
        border(p, "bottom", size=8, color="C62A26", space=6)
        for n, cap in self.figures:
            par = self.doc.add_paragraph()
            par.paragraph_format.space_after = Pt(3)
            par.paragraph_format.tab_stops.add_tab_stop(CONTENT_W, WD_TAB_ALIGNMENT.RIGHT, 2)
            style_run(par.add_run("Figure %d   " % n), font=SANS, size=9.5, color=RED, bold=True)
            style_run(par.add_run(cap), font=SERIF, size=9.5, color=INK)
            style_run(par.add_run("\t"), font=SERIF, size=9.5)
            style_run(par.add_run(str(self.fig_pages.get(n, ""))), font=MONO, size=9, color=INK2)

    def footer(self):
        for sec in self.doc.sections:
            sec.different_first_page_header_footer = True
            f = sec.footer.paragraphs[0]
            f.paragraph_format.tab_stops.add_tab_stop(CONTENT_W, WD_TAB_ALIGNMENT.RIGHT)
            style_run(f.add_run("GeoCras · Cahier des charges"), font=SANS, size=8, color=MUTED)
            style_run(f.add_run("\t"), font=SANS, size=8)
            r = field(f, " PAGE ")
            style_run(r, font=MONO, size=9, color=INK, bold=True)
            border(f, "top", size=4, color="D9D4C9", space=6)

    def save(self, path):
        self.footer()
        self.doc.save(path)


# --------------------------------------------------------------------------- #
def precount(blocks):
    """Numérote titres et figures sans rien rendre, pour que le sommaire
    puisse être écrit avant les chapitres qu'il annonce."""
    heads, figs, c, fn = [], [], [0, 0, 0], 0
    for kind, payload in blocks:
        if kind and kind.startswith("h") and kind[1:].isdigit():
            n = int(kind[1])
            if n > 0:
                c[n - 1] += 1
                for k in range(n, 3):
                    c[k] = 0
            if n == 0:
                heads.append((0, "", payload))
            else:
                heads.append((n, ".".join(str(c[k]) for k in range(n)), payload))
        elif kind == "fig":
            fn += 1
            figs.append((fn, payload[1]))
        elif kind == "fig2":
            fn += 1
            figs.append((fn, payload[2]))
    return heads, figs


def heads_before_toc(blocks):
    """Nombre de titres situés avant le sommaire : ils se cherchent en amont."""
    n = 0
    for kind, _ in blocks:
        if kind == "toc":
            return n
        if kind and kind.startswith("h") and kind[1:].isdigit():
            n += 1
    return n


def build(source, out_docx, toc_pages=None, fig_pages=None):
    blocks = parse(source)
    b = Builder(toc_pages, fig_pages)
    b.headings, b.figures = precount(blocks)
    b._pre_heads, b._pre_figs = list(b.headings), list(b.figures)
    b._n_front = heads_before_toc(blocks)
    for kind, payload in blocks:
        if kind == "cover":
            b.cover(payload)
        elif kind == "toc":
            b.render_toc()
        elif kind == "figtoc":
            b.render_figtoc()
        elif kind == "pb":
            p = b.doc.add_paragraph(); page_break_before(p)
        elif kind.startswith("h"):
            b.heading(kind[1], payload)
        elif kind == "p":
            b.para(payload)
        elif kind == "fig":
            b.figure(*payload)
        elif kind == "fig2":
            b.figure_pair(*payload)
        elif kind == "tab":
            b.table(*payload)
        elif kind == "code":
            b.code(*payload)
        elif kind == "note":
            b.note(*payload)
        elif kind == "list":
            b.bullets(payload)
    b.save(out_docx)
    return b


def to_pdf(docx_path, outdir):
    subprocess.run(["soffice", "--headless", "-env:UserInstallation=file:///tmp/lo_build",
                    "--convert-to", "pdf", "--outdir", outdir, docx_path],
                   check=True, capture_output=True, timeout=600)
    return docx_path.replace(".docx", ".pdf")
